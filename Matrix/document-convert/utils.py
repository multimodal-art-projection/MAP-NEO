import argparse
import copy
import os
import tarfile
import time
import traceback
import warnings
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Union

import cv2
import numpy as np
import urllib.request
import yaml
from onnxruntime import (
    GraphOptimizationLevel,
    InferenceSession,
    SessionOptions,
    get_available_providers,
    get_device,
)
from PIL import Image, UnidentifiedImageError

import re
import docx
from docx import Document
from bs4 import BeautifulSoup
from html.parser import HTMLParser

from lxml import html
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from premailer import Premailer

root_dir = Path(__file__).resolve().parent
InputType = Union[str, np.ndarray, bytes, Path]

def sorted_boxes(dt_boxes):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        dt_boxes(array):detected text boxes with shape [4, 2]
    return:
        sorted boxes(array) with shape [4, 2]
    """
    num_boxes = dt_boxes.shape[0]
    sorted_boxes = sorted(dt_boxes, key=lambda x: (x[0][1], x[0][0]))
    _boxes = list(sorted_boxes)

    for i in range(num_boxes - 1):
        for j in range(i, -1, -1):
            if abs(_boxes[j + 1][0][1] - _boxes[j][0][1]) < 10 and \
                    (_boxes[j + 1][0][0] < _boxes[j][0][0]):
                tmp = _boxes[j]
                _boxes[j] = _boxes[j + 1]
                _boxes[j + 1] = tmp
            else:
                break
    return _boxes

class OrtInferSession:
    def __init__(self, config):
        sess_opt = SessionOptions()
        sess_opt.log_severity_level = 4
        sess_opt.enable_cpu_mem_arena = False
        sess_opt.graph_optimization_level = GraphOptimizationLevel.ORT_ENABLE_ALL

        cpu_ep = "CPUExecutionProvider"
        cpu_provider_options = {
            "arena_extend_strategy": "kSameAsRequested",
        }

        cuda_ep = "CUDAExecutionProvider"
        cuda_provider_options = {
            "device_id": 0,
            "arena_extend_strategy": "kNextPowerOfTwo",
            "cudnn_conv_algo_search": "EXHAUSTIVE",
            "do_copy_in_default_stream": True,
        }

        EP_list = []
        if (
            config["use_cuda"]
            and get_device() == "GPU"
            and cuda_ep in get_available_providers()
        ):
            EP_list = [(cuda_ep, cuda_provider_options)]
        EP_list.append((cpu_ep, cpu_provider_options))

        self._verify_model(config["model_path"])
        self.session = InferenceSession(
            config["model_path"], sess_options=sess_opt, providers=EP_list
        )

        if config["use_cuda"] and cuda_ep not in self.session.get_providers():
            warnings.warn(
                f"{cuda_ep} is not avaiable for current env, the inference part is automatically shifted to be executed under {cpu_ep}.\n"
                "Please ensure the installed onnxruntime-gpu version matches your cuda and cudnn version, "
                "you can check their relations from the offical web site: "
                "https://onnxruntime.ai/docs/execution-providers/CUDA-ExecutionProvider.html",
                RuntimeWarning,
            )

    def __call__(self, input_content: np.ndarray) -> np.ndarray:
        input_dict = dict(zip(self.get_input_names(), [input_content]))
        try:
            return self.session.run(self.get_output_names(), input_dict)
        except Exception as e:
            error_info = traceback.format_exc()
            raise ONNXRuntimeError(error_info) from e

    def get_input_names(
        self,
    ):
        return [v.name for v in self.session.get_inputs()]

    def get_output_names(
        self,
    ):
        return [v.name for v in self.session.get_outputs()]

    def get_character_list(self, key: str = "character"):
        return self.meta_dict[key].splitlines()

    def have_key(self, key: str = "character") -> bool:
        self.meta_dict = self.session.get_modelmeta().custom_metadata_map
        if key in self.meta_dict.keys():
            return True
        return False

    @staticmethod
    def _verify_model(model_path):
        model_path = Path(model_path)
        if not model_path.exists():
            raise FileNotFoundError(f"{model_path} does not exists.")
        if not model_path.is_file():
            raise FileExistsError(f"{model_path} is not a file.")


class ONNXRuntimeError(Exception):
    pass

def download_and_extract_models(model_url='https://molar-public.oss-cn-hangzhou.aliyuncs.com/models.tar.gz', models_folder='models'): 
    try:
        if not os.path.exists(models_folder):
            os.makedirs(models_folder)

            download_path = os.path.join("", 'models.tar.gz')

            urllib.request.urlretrieve(model_url, download_path)

            with tarfile.open(download_path, 'r:gz') as tar:
                tar.extractall(path="")

            os.remove(download_path)
    except urllib.error.URLError as e:
        print(f"Error downloading the model: {e}")
    except tarfile.ReadError as e:
        print(f"Error reading or extracting the tar file: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")


class LoadImage:
    def __init__(
        self,
    ):
        pass

    def __call__(self, img: InputType) -> np.ndarray:
        if not isinstance(img, InputType.__args__):
            raise LoadImageError(
                f"The img type {type(img)} does not in {InputType.__args__}"
            )

        img = self.load_img(img)
        img = self.convert_img(img)
        return img

    def load_img(self, img: InputType) -> np.ndarray:
        if isinstance(img, (str, Path)):
            self.verify_exist(img)
            try:
                img = np.array(Image.open(img))
            except UnidentifiedImageError as e:
                raise LoadImageError(f"cannot identify image file {img}") from e
            return img

        if isinstance(img, bytes):
            img = np.array(Image.open(BytesIO(img)))
            return img

        if isinstance(img, np.ndarray):
            return img

        raise LoadImageError(f"{type(img)} is not supported!")

    def convert_img(self, img: np.ndarray):
        if img.ndim == 2:
            return cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)

        if img.ndim == 3:
            channel = img.shape[2]
            if channel == 1:
                return cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)

            if channel == 2:
                return self.cvt_two_to_three(img)

            if channel == 4:
                return self.cvt_four_to_three(img)

            if channel == 3:
                return cv2.cvtColor(img, cv2.COLOR_RGB2BGR)

            raise LoadImageError(
                f"The channel({channel}) of the img is not in [1, 2, 3, 4]"
            )

        raise LoadImageError(f"The ndim({img.ndim}) of the img is not in [2, 3]")

    @staticmethod
    def cvt_four_to_three(img: np.ndarray) -> np.ndarray:
        """RGBA → BGR"""
        r, g, b, a = cv2.split(img)
        new_img = cv2.merge((b, g, r))

        not_a = cv2.bitwise_not(a)
        not_a = cv2.cvtColor(not_a, cv2.COLOR_GRAY2BGR)

        new_img = cv2.bitwise_and(new_img, new_img, mask=a)
        new_img = cv2.add(new_img, not_a)
        return new_img

    @staticmethod
    def cvt_two_to_three(img: np.ndarray) -> np.ndarray:
        """gray + alpha → BGR"""
        img_gray = img[..., 0]
        img_bgr = cv2.cvtColor(img_gray, cv2.COLOR_GRAY2BGR)

        img_alpha = img[..., 1]
        not_a = cv2.bitwise_not(img_alpha)
        not_a = cv2.cvtColor(not_a, cv2.COLOR_GRAY2BGR)

        new_img = cv2.bitwise_and(img_bgr, img_bgr, mask=img_alpha)
        new_img = cv2.add(new_img, not_a)
        return new_img

    @staticmethod
    def verify_exist(file_path: Union[str, Path]):
        if not Path(file_path).exists():
            raise LoadImageError(f"{file_path} does not exist.")


class LoadImageError(Exception):
    pass


def read_yaml(yaml_path):
    with open(yaml_path, "rb") as f:
        data = yaml.load(f, Loader=yaml.Loader)
    return data


def concat_model_path(config):
    key = "model_path"
    config["Det"][key] = str(root_dir / config["Det"][key])
    config["Rec"][key] = str(root_dir / config["Rec"][key])
    config["Cls"][key] = str(root_dir / config["Cls"][key])
    config["Struc"][key] = str(root_dir / config["Struc"][key])
    config["Table"][key] = str(root_dir / config["Table"][key])
    return config


def init_args():
    parser = argparse.ArgumentParser()
    parser.add_argument("-img", "--img_path", type=str, default=None, required=True)
    parser.add_argument("-p", "--print_cost", action="store_true", default=False)

    global_group = parser.add_argument_group(title="Global")
    global_group.add_argument("--text_score", type=float, default=0.5)
    global_group.add_argument("--use_angle_cls", type=bool, default=True)
    global_group.add_argument("--use_text_det", type=bool, default=True)
    global_group.add_argument("--print_verbose", type=bool, default=False)
    global_group.add_argument("--min_height", type=int, default=30)
    global_group.add_argument("--width_height_ratio", type=int, default=8)

    det_group = parser.add_argument_group(title="Det")
    det_group.add_argument("--det_use_cuda", action="store_true", default=False)
    det_group.add_argument("--det_model_path", type=str, default=None)
    det_group.add_argument("--det_limit_side_len", type=float, default=736)
    det_group.add_argument(
        "--det_limit_type", type=str, default="min", choices=["max", "min"]
    )
    det_group.add_argument("--det_thresh", type=float, default=0.3)
    det_group.add_argument("--det_box_thresh", type=float, default=0.5)
    det_group.add_argument("--det_unclip_ratio", type=float, default=1.6)
    det_group.add_argument("--det_use_dilation", type=bool, default=True)
    det_group.add_argument(
        "--det_score_mode", type=str, default="fast", choices=["slow", "fast"]
    )

    cls_group = parser.add_argument_group(title="Cls")
    cls_group.add_argument("--cls_use_cuda", action="store_true", default=False)
    cls_group.add_argument("--cls_model_path", type=str, default=None)
    cls_group.add_argument("--cls_image_shape", type=list, default=[3, 48, 192])
    cls_group.add_argument("--cls_label_list", type=list, default=["0", "180"])
    cls_group.add_argument("--cls_batch_num", type=int, default=6)
    cls_group.add_argument("--cls_thresh", type=float, default=0.9)

    rec_group = parser.add_argument_group(title="Rec")
    rec_group.add_argument("--rec_use_cuda", action="store_true", default=False)
    rec_group.add_argument("--rec_model_path", type=str, default=None)
    rec_group.add_argument("--rec_img_shape", type=list, default=[3, 48, 320])
    rec_group.add_argument("--rec_batch_num", type=int, default=6)

    args = parser.parse_args()
    return args


class UpdateParameters:
    def __init__(self) -> None:
        pass

    def parse_kwargs(self, **kwargs):
        global_dict, det_dict, cls_dict, rec_dict = {}, {}, {}, {}
        for k, v in kwargs.items():
            if k.startswith("det"):
                det_dict[k] = v
            elif k.startswith("cls"):
                cls_dict[k] = v
            elif k.startswith("rec"):
                rec_dict[k] = v
            else:
                global_dict[k] = v
        return global_dict, det_dict, cls_dict, rec_dict

    def __call__(self, config, **kwargs):
        global_dict, det_dict, cls_dict, rec_dict = self.parse_kwargs(**kwargs)
        new_config = {
            "Global": self.update_global_params(config["Global"], global_dict),
            "Det": self.update_det_params(config["Det"], det_dict),
            "Cls": self.update_cls_params(config["Cls"], cls_dict),
            "Rec": self.update_rec_params(config["Rec"], rec_dict),
        }
        return new_config

    def update_global_params(self, config, global_dict):
        if global_dict:
            config.update(global_dict)
        return config

    def update_det_params(self, config, det_dict):
        if not det_dict:
            return config

        det_dict = {k.split("det_")[1]: v for k, v in det_dict.items()}
        model_path = det_dict.get("model_path", None)
        if not model_path:
            det_dict["model_path"] = str(root_dir / config["model_path"])

        config.update(det_dict)
        return config

    def update_cls_params(self, config, cls_dict):
        if not cls_dict:
            return config

        need_remove_prefix = ["cls_label_list", "cls_model_path", "cls_use_cuda"]
        new_cls_dict = self.remove_prefix(cls_dict, "cls_", need_remove_prefix)

        model_path = new_cls_dict.get("model_path", None)
        if model_path:
            new_cls_dict["model_path"] = str(root_dir / config["model_path"])

        config.update(new_cls_dict)
        return config

    def update_rec_params(self, config, rec_dict):
        if not rec_dict:
            return config

        need_remove_prefix = ["rec_model_path", "rec_use_cuda"]
        new_rec_dict = self.remove_prefix(rec_dict, "rec_", need_remove_prefix)

        model_path = new_rec_dict.get("model_path", None)
        if not model_path:
            new_rec_dict["model_path"] = str(root_dir / config["model_path"])

        config.update(new_rec_dict)
        return config

    @staticmethod
    def remove_prefix(
        config: Dict[str, str], prefix: str, remove_params: List[str]
    ) -> Dict[str, str]:
        new_rec_dict = {}
        for k, v in config.items():
            if k in remove_params:
                k = k.split(prefix)[1]
            new_rec_dict[k] = v
        return new_rec_dict

def string_to_int(s):
    if s.isdigit():
        return int(s)
    return 0


from openpyxl.cell import cell
from openpyxl.styles import Font, Alignment, PatternFill, NamedStyle, Border, Side, Color
from openpyxl.styles.fills import FILL_SOLID
from openpyxl.styles.numbers import FORMAT_CURRENCY_USD_SIMPLE, FORMAT_PERCENTAGE
from openpyxl.styles.colors import BLACK

FORMAT_DATE_MMDDYYYY = 'mm/dd/yyyy'


def colormap(color):
    """
    Convenience for looking up known colors
    """
    cmap = {'black': BLACK}
    return cmap.get(color, color)


def style_string_to_dict(style):
    """
    Convert css style string to a python dictionary
    """
    def clean_split(string, delim):
        return (s.strip() for s in string.split(delim))
    styles = [clean_split(s, ":") for s in style.split(";") if ":" in s]
    return dict(styles)


def get_side(style, name):
    return {'border_style': style.get('border-{}-style'.format(name)),
            'color': colormap(style.get('border-{}-color'.format(name)))}

known_styles = {}


def style_dict_to_named_style(style_dict, number_format=None):
    """
    Change css style (stored in a python dictionary) to openpyxl NamedStyle
    """

    style_and_format_string = str({
        'style_dict': style_dict,
        'parent': style_dict.parent,
        'number_format': number_format,
    })

    if style_and_format_string not in known_styles:
        # Font
        font = Font(bold=style_dict.get('font-weight') == 'bold',
                    color=style_dict.get_color('color', None),
                    size=style_dict.get('font-size'))

        # Alignment
        alignment = Alignment(horizontal=style_dict.get('text-align', 'general'),
                              vertical=style_dict.get('vertical-align'),
                              wrap_text=style_dict.get('white-space', 'nowrap') == 'normal')

        # Fill
        bg_color = style_dict.get_color('background-color')
        fg_color = style_dict.get_color('foreground-color', Color())
        fill_type = style_dict.get('fill-type')
        if bg_color and bg_color != 'transparent':
            fill = PatternFill(fill_type=fill_type or FILL_SOLID,
                               start_color=bg_color,
                               end_color=fg_color)
        else:
            fill = PatternFill()

        # Border
        border = Border(left=Side(**get_side(style_dict, 'left')),
                        right=Side(**get_side(style_dict, 'right')),
                        top=Side(**get_side(style_dict, 'top')),
                        bottom=Side(**get_side(style_dict, 'bottom')),
                        diagonal=Side(**get_side(style_dict, 'diagonal')),
                        diagonal_direction=None,
                        outline=Side(**get_side(style_dict, 'outline')),
                        vertical=None,
                        horizontal=None)

        name = 'Style {}'.format(len(known_styles) + 1)

        pyxl_style = NamedStyle(name=name, font=font, fill=fill, alignment=alignment, border=border,
                                number_format=number_format)

        known_styles[style_and_format_string] = pyxl_style

    return known_styles[style_and_format_string]


class StyleDict(dict):
    """
    It's like a dictionary, but it looks for items in the parent dictionary
    """
    def __init__(self, *args, **kwargs):
        self.parent = kwargs.pop('parent', None)
        super(StyleDict, self).__init__(*args, **kwargs)

    def __getitem__(self, item):
        if item in self:
            return super(StyleDict, self).__getitem__(item)
        elif self.parent:
            return self.parent[item]
        else:
            raise KeyError('{} not found'.format(item))

    def __hash__(self):
        return hash(tuple([(k, self.get(k)) for k in self._keys()]))

    # Yielding the keys avoids creating unnecessary data structures
    # and happily works with both python2 and python3 where the
    # .keys() method is a dictionary_view in python3 and a list in python2.
    def _keys(self):
        yielded = set()
        for k in self.keys():
            yielded.add(k)
            yield k
        if self.parent:
            for k in self.parent._keys():
                if k not in yielded:
                    yielded.add(k)
                    yield k

    def get(self, k, d=None):
        try:
            return self[k]
        except KeyError:
            return d

    def get_color(self, k, d=None):
        """
        Strip leading # off colors if necessary
        """
        color = self.get(k, d)
        if hasattr(color, 'startswith') and color.startswith('#'):
            color = color[1:]
            if len(color) == 3:  # Premailers reduces colors like #00ff00 to #0f0, openpyxl doesn't like that
                color = ''.join(2 * c for c in color)
        return color



class Element(object):
    """
    Our base class for representing an html element along with a cascading style.
    The element is created along with a parent so that the StyleDict that we store
    can point to the parent's StyleDict.
    """
    def __init__(self, element, parent=None):
        self.element = element
        self.number_format = None
        parent_style = parent.style_dict if parent else None
        self.style_dict = StyleDict(style_string_to_dict(element.get('style', '')), parent=parent_style)
        self._style_cache = None

    def style(self):
        """
        Turn the css styles for this element into an openpyxl NamedStyle.
        """
        if not self._style_cache:
            self._style_cache = style_dict_to_named_style(self.style_dict, number_format=self.number_format)
        return self._style_cache

    def get_dimension(self, dimension_key):
        """
        Extracts the dimension from the style dict of the Element and returns it as a float.
        """
        dimension = self.style_dict.get(dimension_key)
        if dimension:
            if dimension[-2:] in ['px', 'em', 'pt', 'in', 'cm']:
                dimension = dimension[:-2]
            dimension = float(dimension)
        return dimension


class TableHead(Element):
    """
    This class maps to the `<th>` element of the html table.
    """
    def __init__(self, head, parent=None):
        super(TableHead, self).__init__(head, parent=parent)
        self.rows = [TableRow(tr, parent=self) for tr in head.findall('tr')]


class TableBody(Element):
    """
    This class maps to the `<tbody>` element of the html table.
    """
    def __init__(self, body, parent=None):
        super(TableBody, self).__init__(body, parent=parent)
        self.rows = [TableRow(tr, parent=self) for tr in body.findall('tr')]


class TableRow(Element):
    """
    This class maps to the `<tr>` element of the html table.
    """
    def __init__(self, tr, parent=None):
        super(TableRow, self).__init__(tr, parent=parent)
        self.cells = [TableCell(cell, parent=self) for cell in tr.findall('th') + tr.findall('td')]



class Table(Element):
    """
    The concrete implementations of Elements are semantically named for the types of elements we are interested in.
    This defines a very concrete tree structure for html tables that we expect to deal with. I prefer this compared to
    allowing Element to have an arbitrary number of children and dealing with an abstract element tree.
    """
    def __init__(self, table):
        """
        takes an html table object (from lxml)
        """
        super(Table, self).__init__(table)
        table_head = table.find('thead')
        self.head = TableHead(table_head, parent=self) if table_head is not None else None
        table_body = table.find('tbody')
        self.body = TableBody(table_body if table_body is not None else table, parent=self)


def get_Tables(doc):
    tree = html.fromstring(doc)
    comments = tree.xpath('//comment()')
    for comment in comments:
        comment.drop_tag()
    return [Table(table) for table in tree.xpath('//table')]


def write_rows(worksheet, elem, row, column=1):
    """
    Writes every tr child element of elem to a row in the worksheet
    returns the next row after all rows are written
    """
    from openpyxl.cell.cell import MergedCell

    initial_column = column
    for table_row in elem.rows:
        for table_cell in table_row.cells:
            cell = worksheet.cell(row=row, column=column)
            while isinstance(cell, MergedCell):
                column += 1
                cell = worksheet.cell(row=row, column=column)

            colspan = string_to_int(table_cell.element.get("colspan", "1"))
            rowspan = string_to_int(table_cell.element.get("rowspan", "1"))
            if rowspan > 1 or colspan > 1:
                worksheet.merge_cells(start_row=row, start_column=column,
                                      end_row=row + rowspan - 1, end_column=column + colspan - 1)

            cell.value = table_cell.value
            table_cell.format(cell)
            min_width = table_cell.get_dimension('min-width')
            max_width = table_cell.get_dimension('max-width')

            if colspan == 1:
                # Initially, when iterating for the first time through the loop, the width of all the cells is None.
                # As we start filling in contents, the initial width of the cell (which can be retrieved by:
                # worksheet.column_dimensions[get_column_letter(column)].width) is equal to the width of the previous
                # cell in the same column (i.e. width of A2 = width of A1)
                width = max(worksheet.column_dimensions[get_column_letter(column)].width or 0, len(table_cell.value) + 2)
                if max_width and width > max_width:
                    width = max_width
                elif min_width and width < min_width:
                    width = min_width
                worksheet.column_dimensions[get_column_letter(column)].width = width
            column += colspan
        row += 1
        column = initial_column
    return row


def table_to_sheet(table, wb):
    """
    Takes a table and workbook and writes the table to a new sheet.
    The sheet title will be the same as the table attribute name.
    """
    ws = wb.create_sheet(title=table.element.get('name'))
    insert_table(table, ws, 1, 1)


def document_to_workbook(doc, wb=None, base_url=None):
    """
    Takes a string representation of an html document and writes one sheet for
    every table in the document.
    The workbook is returned
    """
    if not wb:
        wb = Workbook()
        wb.remove(wb.active)

    inline_styles_doc = Premailer(doc, base_url=base_url, remove_classes=False).transform()
    tables = get_Tables(inline_styles_doc)

    for table in tables:
        table_to_sheet(table, wb)

    return wb


def document_to_xl(doc, filename, base_url=None):
    """
    Takes a string representation of an html document and writes one sheet for
    every table in the document. The workbook is written out to a file called filename
    """
    wb = document_to_workbook(doc, base_url=base_url)
    wb.save(filename)


def insert_table(table, worksheet, column, row):
    if table.head:
        row = write_rows(worksheet, table.head, row, column)
    if table.body:
        row = write_rows(worksheet, table.body, row, column)


def insert_table_at_cell(table, cell):
    """
    Inserts a table at the location of an openpyxl Cell object.
    """
    ws = cell.parent
    column, row = cell.column, cell.row
    insert_table(table, ws, column, row)

def _element_to_string(el):
    string = ''

    for x in el.iterchildren():
        string += '\n' + _element_to_string(x)

    text = el.text.strip() if el.text else ''
    tail = el.tail.strip() if el.tail else ''

    return text + string + '\n' + tail

def element_to_string(el):
    return _element_to_string(el).strip()

class TableCell(Element):
    """
    This class maps to the `<td>` element of the html table.
    """
    CELL_TYPES = {'TYPE_STRING', 'TYPE_FORMULA', 'TYPE_NUMERIC', 'TYPE_BOOL', 'TYPE_CURRENCY', 'TYPE_PERCENTAGE',
                  'TYPE_NULL', 'TYPE_INLINE', 'TYPE_ERROR', 'TYPE_FORMULA_CACHE_STRING', 'TYPE_INTEGER'}

    def __init__(self, cell, parent=None):
        super(TableCell, self).__init__(cell, parent=parent)
        self.value = element_to_string(cell)
        self.number_format = self.get_number_format()

    def data_type(self):
        cell_types = self.CELL_TYPES & set(self.element.get('class', '').split())
        if cell_types:
            if 'TYPE_FORMULA' in cell_types:
                # Make sure TYPE_FORMULA takes precedence over the other classes in the set.
                cell_type = 'TYPE_FORMULA'
            elif cell_types & {'TYPE_CURRENCY', 'TYPE_INTEGER', 'TYPE_PERCENTAGE'}:
                cell_type = 'TYPE_NUMERIC'
            else:
                cell_type = cell_types.pop()
        else:
            cell_type = 'TYPE_STRING'
        return getattr(cell, cell_type)

    def get_number_format(self):
        if 'TYPE_CURRENCY' in self.element.get('class', '').split():
            return FORMAT_CURRENCY_USD_SIMPLE
        if 'TYPE_INTEGER' in self.element.get('class', '').split():
            return '#,##0'
        if 'TYPE_PERCENTAGE' in self.element.get('class', '').split():
            return FORMAT_PERCENTAGE
        if 'TYPE_DATE' in self.element.get('class', '').split():
            return FORMAT_DATE_MMDDYYYY
        if self.data_type() == cell.TYPE_NUMERIC:
            try:
                int(self.value)
            except ValueError:
                return '#,##0.##'
            else:
                return '#,##0'

    def format(self, cell):
        cell.style = self.style()
        data_type = self.data_type()
        if data_type:
            cell.data_type = data_type


def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]['layout'] = 'single'
        return res

    sorted_boxes = sorted(res, key=lambda x: (x['bbox'][1], x['bbox'][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if _boxes[i]['bbox'][1] > _boxes[i - 1]['bbox'][3] and _boxes[i][
                    'bbox'][0] < w / 2 and _boxes[i]['bbox'][2] > w / 2:
                new_res += res_left
                new_res += res_right
                _boxes[i]['layout'] = 'single'
                new_res.append(_boxes[i])
            else:
                if _boxes[i]['bbox'][2] > w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]['bbox'][0] < w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]['bbox'][0] < w / 4 and _boxes[i]['bbox'][2] < 3 * w / 4:
            _boxes[i]['layout'] = 'double'
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]['bbox'][0] > w / 4 and _boxes[i]['bbox'][2] > w / 2:
            _boxes[i]['layout'] = 'double'
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]['layout'] = 'single'
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res



from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def convert_info_docx(img, res, save_folder, img_name):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)

    flag = 1
    for i, region in enumerate(res):
        img_idx = region['img_idx']
        if flag == 2 and region['layout'] == 'single':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
            flag = 1
        elif flag == 1 and region['layout'] == 'double':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
            flag = 2

        if region['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder,
                                    '{}_{}.jpg'.format(region['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif region['type'].lower() == 'title':
            try:
                doc.add_heading(region['res'][0]['text'])
            except Exception as error:
                print(region['res'])
                print(error)
        elif region['type'].lower() == 'table':
            try:
                if region['res']['html'] is None:
                    excel_save_folder = os.path.join(save_folder, img_name)
                    img_path = os.path.join(excel_save_folder,
                                            '{}_{}.jpg'.format(region['bbox'], img_idx))
                    paragraph_pic = doc.add_paragraph()
                    paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph_pic.add_run("")
                    if flag == 1:
                        run.add_picture(img_path, width=shared.Inches(5))
                    elif flag == 2:
                        run.add_picture(img_path, width=shared.Inches(2))
                else:
                    parser = HtmlToDocx()
                    parser.table_style = 'TableGrid'
                    parser.handle_table(region['res']['html'], doc)
            except Exception as error:
                print(region['res'])
                print(error)
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for i, line in enumerate(region['res']):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line['text'] + ' ')
                text_run.font.size = shared.Pt(10)

    # save to docx
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    doc.save(docx_path)

def get_table_rows(table_soup):
    table_row_selectors = [
        'table > tr', 'table > thead > tr', 'table > tbody > tr',
        'table > tfoot > tr'
    ]
    # If there's a header, body, footer or direct child tr tags, add row dimensions from there
    return table_soup.select(', '.join(table_row_selectors), recursive=False)


def get_table_columns(row):
    # Get all columns for the specified row tag.
    return row.find_all(['th', 'td'], recursive=False) if row else []


def get_table_dimensions(table_soup):
    # Get rows for the table
    rows = get_table_rows(table_soup)
    # Table is either empty or has non-direct children between table and tr tags
    # Thus the row dimensions and column dimensions are assumed to be 0

    cols = get_table_columns(rows[0]) if rows else []
    # Add colspan calculation column number
    col_count = 0
    for col in cols:
        colspan = col.attrs.get('colspan', 1)
        col_count += int(colspan)

    return rows, col_count


def get_cell_html(soup):
    # Returns string of td element with opening and closing <td> tags removed
    # Cannot use find_all as it only finds element tags and does not find text which
    # is not inside an element
    return ' '.join([str(i) for i in soup.contents])


def delete_paragraph(paragraph):
    # https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def remove_whitespace(string, leading=False, trailing=False):
    """Remove white space from a string.
    Args:
        string(str): The string to remove white space from.
        leading(bool, optional): Remove leading new lines when True.
        trailing(bool, optional): Remove trailing new lines when False.
    Returns:
        str: The input string with new line characters removed and white space squashed.
    Examples:
        Single or multiple new line characters are replaced with space.
            >>> remove_whitespace("abc\\ndef")
            'abc def'
            >>> remove_whitespace("abc\\n\\n\\ndef")
            'abc def'
        New line characters surrounded by white space are replaced with a single space.
            >>> remove_whitespace("abc \\n \\n \\n def")
            'abc def'
            >>> remove_whitespace("abc  \\n  \\n  \\n  def")
            'abc def'
        Leading and trailing new lines are replaced with a single space.
            >>> remove_whitespace("\\nabc")
            ' abc'
            >>> remove_whitespace("  \\n  abc")
            ' abc'
            >>> remove_whitespace("abc\\n")
            'abc '
            >>> remove_whitespace("abc  \\n  ")
            'abc '
        Use ``leading=True`` to remove leading new line characters, including any surrounding
        white space:
            >>> remove_whitespace("\\nabc", leading=True)
            'abc'
            >>> remove_whitespace("  \\n  abc", leading=True)
            'abc'
        Use ``trailing=True`` to remove trailing new line characters, including any surrounding
        white space:
            >>> remove_whitespace("abc  \\n  ", trailing=True)
            'abc'
    """
    # Remove any leading new line characters along with any surrounding white space
    if leading:
        string = re.sub(r'^\s*\n+\s*', '', string)

    # Remove any trailing new line characters along with any surrounding white space
    if trailing:
        string = re.sub(r'\s*\n+\s*$', '', string)

    # Replace new line characters and absorb any surrounding space.
    string = re.sub(r'\s*\n\s*', ' ', string)
    # TODO need some way to get rid of extra spaces in e.g. text <span>   </span>  text
    return re.sub(r'\s+', ' ', string)


font_styles = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
}

font_names = {
    'code': 'Courier',
    'pre': 'Courier',
}


class HtmlToDocx(HTMLParser):
    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'styles': True,
        }
        self.table_row_selectors = [
            'table > tr', 'table > thead > tr', 'table > tbody > tr',
            'table > tfoot > tr'
        ]
        self.table_style = None
        self.paragraph_style = None

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options[
            'fix-html']  # whether or not to clean with BeautifulSoup
        self.document = self.doc
        self.include_tables = True  #TODO add this option back in?
        self.include_images = self.options['images']
        self.include_styles = self.options['styles']
        self.paragraph = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0

    def copy_settings_from(self, other):
        """Copy settings from another instance of HtmlToDocx"""
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def ignore_nested_tables(self, tables_soup):
        """
        Returns array containing only the highest level tables
        Operates on the assumption that bs4 returns child elements immediately after
        the parent element in `find_all`. If this changes in the future, this method will need to be updated
        :return:
        """
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables

    def get_tables(self):
        if not hasattr(self, 'soup'):
            self.include_tables = False
            return
            # find other way to do it, or require this dependency?
        self.tables = self.ignore_nested_tables(self.soup.find_all('table'))
        self.table_no = 0

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            html = str(self.soup)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' %
                             docx.table._Cell)
        unwanted_paragraph = cell.paragraphs[0]
        if unwanted_paragraph.text == "":
            delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        # cells must end with a paragraph or will get message about corrupt file
        # https://stackoverflow.com/a/29287121
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')

    def apply_paragraph_style(self, style=None):
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(
                f"Unable to apply style {self.paragraph_style}.") from e

    def handle_table(self, html, doc):
        """
        To handle nested tables, we will parse tables manually as follows:
        Get table soup
        Create docx table
        Iterate over soup and fill docx table with new instances of this parser
        Tell HTMLParser to ignore any tags until the corresponding closing table tag
        """
        table_soup = BeautifulSoup(html, 'html.parser')
        rows, cols_len = get_table_dimensions(table_soup)
        table = doc.add_table(len(rows), cols_len)
        table.style = doc.styles['Table Grid']

        cell_row = 0
        for index, row in enumerate(rows):
            cols = get_table_columns(row)
            cell_col = 0
            for col in cols:
                colspan = int(col.attrs.get('colspan', 1))
                rowspan = int(col.attrs.get('rowspan', 1))

                cell_html = get_cell_html(col)
                if col.name == 'th':
                    cell_html = "<b>%s</b>" % cell_html

                docx_cell = table.cell(cell_row, cell_col)

                while docx_cell.text != '':  # Skip the merged cell
                    cell_col += 1
                    docx_cell = table.cell(cell_row, cell_col)

                cell_to_merge = table.cell(cell_row + rowspan - 1,
                                           cell_col + colspan - 1)
                if docx_cell != cell_to_merge:
                    docx_cell.merge(cell_to_merge)

                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)
                child_parser.add_html_to_cell(cell_html or ' ', docx_cell)

                cell_col += colspan
            cell_row += 1

    def handle_data(self, data):
        if self.skip:
            return

        # Only remove white space if we're not in a pre block.
        if 'pre' not in self.tags:
            # remove leading and trailing whitespace in all instances
            data = remove_whitespace(data, True, True)

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()
            self.apply_paragraph_style()

        # There can only be one nested link in a valid html document
        # You cannot have interactive content in an A tag, this includes links
        # https://html.spec.whatwg.org/#interactive-content
        link = self.tags.get('a')
        if link:
            self.handle_link(link['href'], data)
        else:
            # If there's a link, dont put the data directly in the run
            self.run = self.paragraph.add_run(data)
            spans = self.tags['span']
            for span in spans:
                if 'style' in span:
                    style = self.parse_dict_string(span['style'])
                    self.add_styles_to_run(style)

            # add font style and name
            for tag in self.tags:
                if tag in font_styles:
                    font_style = font_styles[tag]
                    setattr(self.run.font, font_style, True)

                if tag in font_names:
                    font_name = font_names[tag]
                    self.run.font.name = font_name

def get_rotate_crop_image(img, points):
    '''
    img_height, img_width = img.shape[0:2]
    left = int(np.min(points[:, 0]))
    right = int(np.max(points[:, 0]))
    top = int(np.min(points[:, 1]))
    bottom = int(np.max(points[:, 1]))
    img_crop = img[top:bottom, left:right, :].copy()
    points[:, 0] = points[:, 0] - left
    points[:, 1] = points[:, 1] - top
    '''
    assert len(points) == 4, "shape of points must be 4*2"
    img_crop_width = int(
        max(
            np.linalg.norm(points[0] - points[1]),
            np.linalg.norm(points[2] - points[3])))
    img_crop_height = int(
        max(
            np.linalg.norm(points[0] - points[3]),
            np.linalg.norm(points[1] - points[2])))
    pts_std = np.float32([[0, 0], [img_crop_width, 0],
                          [img_crop_width, img_crop_height],
                          [0, img_crop_height]])
    M = cv2.getPerspectiveTransform(points, pts_std)
    dst_img = cv2.warpPerspective(
        img,
        M, (img_crop_width, img_crop_height),
        borderMode=cv2.BORDER_REPLICATE,
        flags=cv2.INTER_CUBIC)
    dst_img_height, dst_img_width = dst_img.shape[0:2]
    if dst_img_height * 1.0 / dst_img_width >= 1.5:
        dst_img = np.rot90(dst_img)
    return dst_img


def get_minarea_rect_crop(img, points):
    bounding_box = cv2.minAreaRect(np.array(points).astype(np.int32))
    points = sorted(list(cv2.boxPoints(bounding_box)), key=lambda x: x[0])

    index_a, index_b, index_c, index_d = 0, 1, 2, 3
    if points[1][1] > points[0][1]:
        index_a = 0
        index_d = 1
    else:
        index_a = 1
        index_d = 0
    if points[3][1] > points[2][1]:
        index_b = 2
        index_c = 3
    else:
        index_b = 3
        index_c = 2

    box = [points[index_a], points[index_b], points[index_c], points[index_d]]
    crop_img = get_rotate_crop_image(img, np.array(box))
    return crop_img



class aggr():
    def __init__(self, det, rec, drop_score = 0.5) -> None:
        self.text_detector = det
        self.text_recognizer = rec
        self.drop_score = drop_score

    def __call__(self, img, cls=True):
        time_dict = {'det': 0, 'rec': 0, 'csl': 0, 'all': 0}
        start = time.time()
        ori_im = img.copy()
        dt_boxes, elapse = self.text_detector(img)
        time_dict['det'] = elapse
        if dt_boxes is None:
            return None, None
        img_crop_list = []

        dt_boxes = sorted_boxes(dt_boxes)

        for bno in range(len(dt_boxes)):
            tmp_box = copy.deepcopy(dt_boxes[bno])
            img_crop = get_rotate_crop_image(ori_im, tmp_box)
            img_crop_list.append(img_crop)

        rec_res, elapse = self.text_recognizer(img_crop_list)
        time_dict['rec'] = elapse
        filter_boxes, filter_rec_res = [], []
        for box, rec_result in zip(dt_boxes, rec_res):
            text, score = rec_result
            if score >= self.drop_score:
                filter_boxes.append(box)
                filter_rec_res.append(rec_result)
        end = time.time()
        time_dict['all'] = end - start
        return filter_boxes, filter_rec_res, time_dict
    
class MarkdownDocument:
    def __init__(self):
        self.content = ""

    def add_figure(self, figure=None):
        if figure:
            # self.content.append(f"![]({figure})\n\n")
            self.content += f"![]({figure})\n\n"
        else:
            print("No figure URL provided, the figure will not be added.")

    def add_table(self, table=None):
        self.content += f"{table}\n\n"
    
    def add_title(self, title, level=1):
        self.content += f"{'#' * level} {title}\n\n"

    def add_text(self, text):
        self.content += f"{text}\n\n"

    def add_bullet_list(self, items, to_real_list=False):
        if to_real_list: # 要不要真加“-”转成markdown的list
            for item in items:
                self.content += f"- {item}\n"
        else:
            for item in items:
                self.content += f"{item}\n"
            self.content += "\n"
        self.content += "\n"
    def get_markdown(self):
        return self.content
    
    def save(self, path):
        with open(path, 'w') as f:
            f.write(self.get_markdown())
        return


def convert_info_md(img, res, save_folder, img_name):
    md_out = MarkdownDocument()
    for i, region in enumerate(res):
        if region['type'].lower() == 'figure':
            try:
                fig_path = os.path.join(img_name, f"{region['bbox']}_0.jpg")
            except:
                fig_path = ""
            md_out.add_figure(fig_path)
        elif region['type'].lower() == 'table':
            md_out.add_table(region['res']['html'])
        elif region['type'].lower() == 'text':
            txt = '\n'.join(it['text'] for it in region['res'])
            md_out.add_text(txt)
        elif region['type'].lower() == 'title':
            txt = ' '.join(it['text'] for it in region['res'])
            md_out.add_title(txt)
        elif region['type'].lower() == 'list':
            txt_list = [it['text'] for it in region['res']]
            md_out.add_bullet_list(txt_list)
        # else:
            
    md_path = os.path.join(save_folder, f'{img_name}_ocr.md')
    md_out.save(md_path)

from copy import deepcopy
import json
import os
import time
import fitz

def read_image(image_file) -> list:
    if os.path.basename(image_file)[-3:] == 'pdf':
        imgs = []
        with fitz.open(image_file) as pdf:
            for pg in range(0, pdf.page_count):
                page = pdf[pg]
                mat = fitz.Matrix(2, 2)
                pm = page.get_pixmap(matrix=mat, alpha=False)

                # if width or height > 2000 pixels, don't enlarge the image
                if pm.width > 2000 or pm.height > 2000:
                    pm = page.get_pixmap(matrix=fitz.Matrix(1, 1), alpha=False)

                img = Image.frombytes("RGB", [pm.width, pm.height], pm.samples)
                img = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                imgs.append(img)
    else:
        img = cv2.imread(image_file, cv2.IMREAD_COLOR)
        if img is not None:
            imgs = [img]

    return imgs


def save_structure_res(res, save_folder, img_name, img_idx=0):
    excel_save_folder = os.path.join(save_folder, img_name)
    os.makedirs(excel_save_folder, exist_ok=True)
    res_cp = deepcopy(res)
    # save res
    with open(
            os.path.join(excel_save_folder, 'res_{}.txt'.format(img_idx)),
            'w',
            encoding='utf8') as f:
        for region in res_cp:
            roi_img = region.pop('img')
            try:
                f.write('{}\n'.format(json.dumps(region)))
            except Exception as e:
                print(e)
                print(region)

            if region['type'].lower() == 'table' and len(region[
                    'res']) > 0 and 'html' in region['res']:
                if region['res']['html'] is None:
                    img_path = os.path.join(
                        excel_save_folder,
                        '{}_{}.jpg'.format(region['bbox'], img_idx))
                    cv2.imwrite(img_path, roi_img)
                else:
                    excel_path = os.path.join(
                        excel_save_folder,
                        '{}_{}.xlsx'.format(region['bbox'], img_idx))
                    to_excel(region['res']['html'], excel_path)
            elif region['type'].lower() == 'figure':
                img_path = os.path.join(
                    excel_save_folder,
                    '{}_{}.jpg'.format(region['bbox'], img_idx))
                cv2.imwrite(img_path, roi_img)
                
def to_excel(html_table, excel_path):
    document_to_xl(html_table, excel_path)